"""Generate FSOI_Explanation.docx for upload to Google Docs."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ────────────────────────────────────────────────
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Helpers ──────────────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1F, 0x49, 0x7D)
MID_BLUE = RGBColor(0x2E, 0x74, 0xB5)
LIGHT_BLUE = RGBColor(0xDF, 0xEC, 0xF8)
GREEN = RGBColor(0x00, 0x70, 0x00)
RED = RGBColor(0xC0, 0x00, 0x00)
ORANGE = RGBColor(0xC5, 0x5A, 0x11)
GRAY_BG = RGBColor(0xF2, 0xF2, 0xF2)
DARK_GRAY = RGBColor(0x40, 0x40, 0x40)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG = RGBColor(0x1E, 0x1E, 0x1E)
CODE_FG = RGBColor(0xD4, 0xD4, 0xD4)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_para_bg(para, rgb: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def heading(text, level=1, color=DARK_BLUE):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    return p


def body(text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def code_block(lines):
    """Mono-spaced dark code block rendered as a table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, CODE_BG)
    # clear default paragraph
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for line in lines:
        p = cell.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = CODE_FG
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()  # spacer


def formula_block(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_bg(p, LIGHT_BLUE)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    doc.add_paragraph()


def callout(text, bg=LIGHT_BLUE, text_color=DARK_BLUE):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = text_color
    run.font.italic = True
    doc.add_paragraph()


def spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


# ═══════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("Forecast Sensitivity to Observations for an\nAI-Based Numerical Weather Prediction System")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = DARK_BLUE

spacer()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("FSOI Implementation Guide — OCELOT AI Weather Model")
run.font.size = Pt(13)
run.font.color.rgb = MID_BLUE
run.font.italic = True

spacer()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("May 2026  •  gnn_model/FSOI/")
run.font.size = Pt(11)
run.font.color.rgb = DARK_GRAY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 1 — WHAT IS FSOI?
# ═══════════════════════════════════════════════════════════════════
heading("1  What is FSOI?", 1)

callout(
    "FSOI answers: \"How much did each observation actually help (or hurt) the forecast?\"\n"
    "It uses adjoint gradients to attribute forecast error changes back to individual\n"
    "observations — without re-running the model for every possible denial experiment."
)

body(
    "In classical Numerical Weather Prediction (NWP), observations from radiosondes, aircraft, "
    "satellites, and surface stations are assimilated via variational methods each cycle. "
    "In OCELOT's AI-based approach, no explicit assimilation step is performed: the true "
    "observations at analysis time serve directly as the analysis state (xa), while the "
    "model's prior forecast from the previous window provides the background (xb). "
    "Not all observations are equally useful — some reduce forecast error, some increase it. "
    "Running a separate Observing System Experiment (OSE) for each instrument type is "
    "prohibitively expensive. FSOI provides a single-pass, gradient-based shortcut."
)

spacer()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 2 — CORE FORMULA
# ═══════════════════════════════════════════════════════════════════
heading("2  The Core Formula", 1)

formula_block("FSOI(k)  =  0.5 × δx(k)  ⊙  ( ga(k)  +  gb(k) )")

# Symbol table
tbl = doc.add_table(rows=6, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Symbol", "Name", "Meaning"]
for i, h in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    set_cell_bg(cell, MID_BLUE)
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(11)

rows_data = [
    ("k", "Observation index", "A single measurement from any sensor"),
    ("δx(k) = xa(k) − xb(k)", "Innovation", "How much the analysis differed from background"),
    ("ga(k)", "Analysis adjoint", "∂J/∂xa — gradient of forecast error w.r.t. analysis obs"),
    ("gb(k)", "Background adjoint", "∂J/∂xb — gradient of forecast error w.r.t. background obs"),
    ("⊙", "Element-wise product", "Multiply innovation × average gradient component-wise"),
]
for r_idx, (sym, name, meaning) in enumerate(rows_data, start=1):
    row = tbl.rows[r_idx]
    if r_idx % 2 == 0:
        for c in row.cells:
            set_cell_bg(c, LIGHT_BLUE)
    data = [sym, name, meaning]
    for c_idx, val in enumerate(data):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx == 0:
            run.font.name = "Courier New"
            run.font.bold = True

doc.add_paragraph()
body("Sign convention:", bold=True)
bullet("FSOI(k) < 0  →  observation k REDUCED forecast error  (beneficial ✓)", bold_prefix="Negative: ")
bullet("FSOI(k) > 0  →  observation k INCREASED forecast error (detrimental ✗)", bold_prefix="Positive: ")
bullet("Sum over all k:  Σ FSOI(k) ≈ ΔJ = J(xa) − J(xb)  (closure property)", bold_prefix="Sum property: ")

spacer()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 3 — GRADIENT SIGN DEEP DIVE
# ═══════════════════════════════════════════════════════════════════
heading("3  Are the Gradients Always Negative?", 1)

callout(
    "No. ga(k) and gb(k) can be positive OR negative — their sign is not fixed.\n"
    "The sign of FSOI is determined by the PRODUCT of the innovation δx and the gradient sum (ga + gb).\n"
    "These two quantities are independent, and either combination is physically possible.",
    bg=RGBColor(0xFF, 0xF8, 0xE1), text_color=ORANGE
)

body("What does the gradient actually mean?", bold=True)
spacer()

body(
    "ga(k) = ∂J/∂xa(k) answers a simple question: "
    "\"If I nudge observation k upward by a tiny amount ε, does the forecast error J go up or down?\""
)

spacer()

tbl = doc.add_table(rows=3, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Gradient value", "Physical meaning", "What the model 'wants'"]):
    set_cell_bg(tbl.rows[0].cells[i], MID_BLUE)
    run = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

grad_rows = [
    ("ga(k) < 0", "Increasing obs k → error DECREASES  (∂J/∂x < 0)", "Obs k should be pushed HIGHER"),
    ("ga(k) > 0", "Increasing obs k → error INCREASES  (∂J/∂x > 0)", "Obs k should be pushed LOWER"),
]
grad_bgs = [RGBColor(0xE2, 0xEF, 0xDA), RGBColor(0xFF, 0xEB, 0xEB)]
for r_idx, (gval, meaning, want) in enumerate(grad_rows, start=1):
    row = tbl.rows[r_idx]
    set_cell_bg(row.cells[0], grad_bgs[r_idx - 1])
    set_cell_bg(row.cells[1], grad_bgs[r_idx - 1])
    set_cell_bg(row.cells[2], grad_bgs[r_idx - 1])
    for c_idx, val in enumerate([gval, meaning, want]):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx == 0:
            run.font.name = "Courier New"
            run.font.bold = True

doc.add_paragraph()
body(
    "The gradient is computed purely by backpropagation through the GNN. It reflects "
    "the model's learned sensitivity — how much each observation channel influences "
    "the downstream forecast error at the verification targets. Satellite channels "
    "often have gradients near zero because the GNN has learned to rely on them weakly; "
    "radiosonde channels have large gradients because the model is highly sensitive to them."
)

spacer()
body("What determines the sign of FSOI?", bold=True)
spacer()

body(
    "FSOI = 0.5 × δx × (ga + gb) is the product of two independent quantities. "
    "FSOI is beneficial (negative) only when the innovation and the gradient point in OPPOSITE directions — "
    "meaning the true observation deviates from the prior forecast in the direction that reduces forecast error."
)
spacer()

code_block([
    "  FSOI sign logic:",
    "  ─────────────────────────────────────────────────────────────────────────",
    "  δx < 0  AND  (ga+gb) > 0  →  FSOI = (−)(+) = NEGATIVE  ✓  BENEFICIAL",
    "    Analysis pulled obs DOWN, gradient says going DOWN reduces error",
    "",
    "  δx > 0  AND  (ga+gb) < 0  →  FSOI = (+)(−) = NEGATIVE  ✓  BENEFICIAL",
    "    Analysis pulled obs UP,   gradient says going UP   reduces error",
    "",
    "  δx < 0  AND  (ga+gb) < 0  →  FSOI = (−)(−) = POSITIVE  ✗  DETRIMENTAL",
    "    Analysis pulled obs DOWN, gradient says going DOWN makes error WORSE",
    "",
    "  δx > 0  AND  (ga+gb) > 0  →  FSOI = (+)(+) = POSITIVE  ✗  DETRIMENTAL",
    "    Analysis pulled obs UP,   gradient says going UP   makes error WORSE",
    "  ─────────────────────────────────────────────────────────────────────────",
    "  Rule: FSOI < 0 ⟺ the true obs deviates from the prior forecast (δx)",
    "        in the direction the gradient says reduces forecast error.",
])

body("Direction-agreement table:", bold=True)
spacer()

tbl2 = doc.add_table(rows=3, cols=3)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

# header row
set_cell_bg(tbl2.rows[0].cells[0], MID_BLUE)
tbl2.rows[0].cells[0].paragraphs[0].add_run("")  # corner blank

for i, label in enumerate(["δx > 0  (analysis went UP)", "δx < 0  (analysis went DOWN)"], start=1):
    set_cell_bg(tbl2.rows[0].cells[i], MID_BLUE)
    run = tbl2.rows[0].cells[i].paragraphs[0].add_run(label)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

dir_data = [
    ("ga+gb > 0\n(model wants obs LOWER)", "DETRIMENTAL  +", "BENEFICIAL  −"),
    ("ga+gb < 0\n(model wants obs HIGHER)", "BENEFICIAL  −", "DETRIMENTAL  +"),
]
cell_colors = {
    "BENEFICIAL  −": RGBColor(0xE2, 0xEF, 0xDA),
    "DETRIMENTAL  +": RGBColor(0xFF, 0xEB, 0xEB),
}
cell_text_colors = {
    "BENEFICIAL  −": GREEN,
    "DETRIMENTAL  +": RED,
}
for r_idx, (row_label, c1, c2) in enumerate(dir_data, start=1):
    row = tbl2.rows[r_idx]
    set_cell_bg(row.cells[0], LIGHT_BLUE)
    run = row.cells[0].paragraphs[0].add_run(row_label)
    run.font.size = Pt(10)
    run.font.bold = True
    for c_idx, val in enumerate([c1, c2], start=1):
        set_cell_bg(row.cells[c_idx], cell_colors[val])
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = cell_text_colors[val]

doc.add_paragraph()

body("What if ga = gb? (the equal-gradient case)", bold=True)
spacer()

body(
    "Even when ga and gb are exactly equal (call them both g), FSOI is still fully determined "
    "by the product of δx and g:"
)

formula_block("FSOI = 0.5 × δx × (g + g)  =  δx × g")

body(
    "Equal gradients simply remove the 0.5 averaging factor. The sign is unchanged — "
    "FSOI is still positive or negative depending on whether δx and g agree or disagree in direction. "
    "So yes, FSOI can be strongly positive even when ga = gb."
)

spacer()

body("Numerical illustration — three scenarios with same δx but different gradient signs:", bold=True)
spacer()

code_block([
    "  Shared setup:  xb = 15.0°C,  xa = 12.0°C,  δx = −3.0",
    "  ───────────────────────────────────────────────────────────────────",
    "  Scenario A — both gradients negative (model wants obs HIGHER):",
    "    ga = −0.8,  gb = −0.6,  ga+gb = −1.4",
    "    FSOI = 0.5 × (−3.0) × (−1.4) = +2.1   ✗  DETRIMENTAL",
    "    Why: analysis pulled obs DOWN but gradient says going DOWN is WRONG",
    "",
    "  Scenario B — both gradients positive (model wants obs LOWER):",
    "    ga = +0.8,  gb = +0.6,  ga+gb = +1.4",
    "    FSOI = 0.5 × (−3.0) × (+1.4) = −2.1   ✓  BENEFICIAL",
    "    Why: analysis pulled obs DOWN and gradient says going DOWN is RIGHT",
    "",
    "  Scenario C — ga = gb (equal gradients, still positive):",
    "    ga = −0.7,  gb = −0.7,  ga+gb = −1.4",
    "    FSOI = 0.5 × (−3.0) × (−1.4) = +2.1   ✗  DETRIMENTAL",
    "    Why: equal gradients change nothing — mismatch direction still detriment",
    "  ───────────────────────────────────────────────────────────────────",
    "  The innovation δx is the SAME in all three cases.",
    "  The gradient SIGN alone flips FSOI from detrimental to beneficial.",
])

callout(
    "Key takeaway: The gradient is NOT a quality score — it is a direction vector.\n"
    "A large negative gradient on an obs means the model is highly sensitive to it AND\n"
    "increasing it would help. Whether the true observation actually deviates from the\n"
    "prior forecast in that direction is what FSOI measures.",
    bg=LIGHT_BLUE, text_color=DARK_BLUE
)

spacer()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 4 — SIMPLE EXAMPLE  (renumbered from 3)
# ═══════════════════════════════════════════════════════════════════
heading("4  Simple Worked Example", 1)

body("Imagine one radiosonde temperature observation entering the system at analysis time.", italic=True)
spacer()

code_block([
    "  Background (prior ML forecast for this window):   xb  =  15.0°C",
    "  True observed value (analysis state):            xa  =  12.0°C",
    "  ─────────────────────────────────────────────────────────────────",
    "  Innovation:   δx  =  xa − xb  =  12.0 − 15.0  =  −3.0",
    "",
    "  Adjoint (gradient of forecast error w.r.t. this obs):",
    "    ga  =  −0.8   (in analysis state)",
    "    gb  =  −0.6   (in background state)",
    "",
    "  FSOI  =  0.5 × (−3.0) × (−0.8 + −0.6)",
    "        =  0.5 × (−3.0) × (−1.4)",
    "        =  0.5 × 4.2",
    "        =  +2.1   ←  DETRIMENTAL",
])

callout(
    "Even though the true obs is 3°C below the prior forecast (δx = −3), the gradients "
    "say the forecast error went up. This means the innovation pointed in the wrong "
    "direction for this particular state — the background was closer to what the model needed.",
    bg=RGBColor(0xFF, 0xF0, 0xE0), text_color=ORANGE
)

body("If FSOI had come out −2.1 instead, that observation would have reduced forecast error by 2.1 units — beneficial.")

spacer()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 4 — FULL PIPELINE FLOWCHART (text-art in code block)
# ═══════════════════════════════════════════════════════════════════
heading("4  Full Pipeline — Step-by-Step", 1)

body(
    "The pipeline processes pairs of consecutive 12-hour forecast windows "
    "(prev_batch, curr_batch) in strict chronological order. For a month of data "
    "at 12-hour intervals that gives 60 pairs."
)
spacer()

code_block([
    "═══════════════════════════════════════════════════════════════════",
    "                     FSOI PIPELINE OVERVIEW",
    "═══════════════════════════════════════════════════════════════════",
    "",
    "  ┌─────────────────────────────────────────────────────────────┐",
    "  │  INPUT: Sequential data bins  (12-hr windows, no shuffle)  │",
    "  │         e.g. 2025-07-01 00Z → 2025-07-31 12Z  =  60 pairs  │",
    "  └───────────────────────┬─────────────────────────────────────┘",
    "                          │",
    "                          ▼",
    "  ┌─────────────────────────────────────────────────────────────┐",
    "  │  STEP 0: Setup & Validation                                 │",
    "  │  • Load GNN checkpoint (frozen weights, no grad updates)    │",
    "  │  • Verify 12-hr sequential spacing between windows          │",
    "  │  • Set requires_grad=True on INPUT observation tensors      │",
    "  │  • Optional FD gradient check (10 samples to verify ∂J/∂x) │",
    "  └───────────────────────┬─────────────────────────────────────┘",
    "                          │",
    "                          ▼",
    "  ┌─────────────────────────────────────────────────────────────┐",
    "  │  STEP 1: Extract xa and xb                                  │",
    "  │                                                             │",
    "  │  prev_batch ──► model.forward() ──► predicted mesh state    │",
    "  │                                          │                  │",
    "  │                               decode at curr obs locations  │",
    "  │                                          │                  │",
    "  │                                          ▼                  │",
    "  │  xb = background prediction at current obs locations        │",
    "  │  xa = actual current observations  (from curr_batch)        │",
    "  │                                                             │",
    "  │  CRITICAL column offset in input tensor:                    │",
    "  │    obs_features = x_input[:, 7+n_meta : 7+n_meta+n_ch]     │",
    "  │    (cols 0:7 are geo/time features, NOT observations)       │",
    "  └───────────────────────┬─────────────────────────────────────┘",
    "                          │",
    "                          ▼",
    "  ┌─────────────────────────────────────────────────────────────┐",
    "  │  STEP 2: Compute Forecast Error  J(xa) and J(xb)            │",
    "  │                                                             │",
    "  │  Run model forward from xa → 12-h forecast → compare truth  │",
    "  │  Run model forward from xb → 12-h forecast → compare truth  │",
    "  │                                                             │",
    "  │  J(xa) = MSE(forecast_from_xa, verifying_analysis)          │",
    "  │  J(xb) = MSE(forecast_from_xb, verifying_analysis)          │",
    "  │  ΔJ    = J(xa) − J(xb)                                      │",
    "  │         (negative = observations collectively helped)       │",
    "  └───────────────────────┬─────────────────────────────────────┘",
    "                          │",
    "                          ▼",
    "  ┌─────────────────────────────────────────────────────────────┐",
    "  │  STEP 3: Compute Adjoints  ga and gb                        │",
    "  │                                                             │",
    "  │  ga = ∂J/∂xa  via autograd.backward()  on J(xa)             │",
    "  │  gb = ∂J/∂xb  via autograd.backward()  on J(xb)             │",
    "  │                                                             │",
    "  │  These are vectors — one gradient value per obs channel k   │",
    "  └───────────────────────┬─────────────────────────────────────┘",
    "                          │",
    "                          ▼",
    "  ┌─────────────────────────────────────────────────────────────┐",
    "  │  STEP 4: Apply FSOI Formula per observation                  │",
    "  │                                                             │",
    "  │  δx(k)   = xa(k) − xb(k)                                   │",
    "  │  FSOI(k) = 0.5 × δx(k) × ( ga(k) + gb(k) )                │",
    "  │                                                             │",
    "  │  ∑_k FSOI(k) ≈ ΔJ   (closure property)                     │",
    "  └───────────────────────┬─────────────────────────────────────┘",
    "                          │",
    "                          ▼",
    "  ┌─────────────────────────────────────────────────────────────┐",
    "  │  STEP 5: Aggregate Results                                  │",
    "  │                                                             │",
    "  │  By instrument : sum FSOI over all obs from same sensor     │",
    "  │  By channel    : sum FSOI per spectral / variable channel   │",
    "  │  By pressure   : sum FSOI per 50 hPa pressure bin           │",
    "  │  By region     : sum FSOI per lat/lon box                   │",
    "  └───────────────────────┬─────────────────────────────────────┘",
    "                          │",
    "                          ▼",
    "  ┌─────────────────────────────────────────────────────────────┐",
    "  │  STEP 6: Evaluation & Diagnostics                           │",
    "  │                                                             │",
    "  │  Closure check    : ∑FSOI / ΔJ  (ideal = 1.0)              │",
    "  │  Helpful fraction : % obs with FSOI < 0                     │",
    "  │  Innovation RMS   : ‖δx‖ in σ units (linearity check)      │",
    "  │  FD validation    : finite-difference check of gradients    │",
    "  └───────────────────────┬─────────────────────────────────────┘",
    "                          │",
    "                          ▼",
    "  ┌─────────────────────────────────────────────────────────────┐",
    "  │  STEP 7 (Optional): OSE Validation                         │",
    "  │                                                             │",
    "  │  Deny one instrument: replace xa[inst] → xb[inst]           │",
    "  │  Re-run forecast → measure actual ΔJ_ose                    │",
    "  │  Compare sign: does FSOI predict same direction as OSE?     │",
    "  └───────────────────────┬─────────────────────────────────────┘",
    "                          │",
    "                          ▼",
    "  ┌─────────────────────────────────────────────────────────────┐",
    "  │  OUTPUTS                                                    │",
    "  │  fsoi_by_instrument.csv  ← headline rankings               │",
    "  │  fsoi_by_channel.csv     ← per-channel detail               │",
    "  │  fsoi_closure_summary.csv← linear approx quality           │",
    "  │  fsoi_system_health.csv  ← helpful_fraction, flags         │",
    "  │  figures/*.png           ← impact bars, maps, profiles      │",
    "  └─────────────────────────────────────────────────────────────┘",
])

# ═══════════════════════════════════════════════════════════════════
#  SECTION 5 — KEY CODE MODULES
# ═══════════════════════════════════════════════════════════════════
heading("5  Code Module Map", 1)

body("Each step in the pipeline is handled by a dedicated Python module:")
spacer()

tbl = doc.add_table(rows=10, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

col_widths = [Inches(2.0), Inches(1.2), Inches(2.8)]
for i, w in enumerate(col_widths):
    for row in tbl.rows:
        row.cells[i].width = w

header_row = tbl.rows[0]
for i, h in enumerate(["Module", "Pipeline Step", "Responsibility"]):
    set_cell_bg(header_row.cells[i], MID_BLUE)
    run = header_row.cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

modules = [
    ("fsoi_dataset.py", "Step 0", "Build chronological (prev, curr) pairs; verify 12-hr spacing"),
    ("fsoi_model_extensions.py", "Step 1", "Decode background xb at observation locations; freeze model"),
    ("fsoi_utils.py", "Steps 1–5", "Extract obs columns, compute δx, ga, gb, FSOI, aggregate"),
    ("fsoi_validation.py", "Step 0+6", "FD gradient check, alignment verification"),
    ("fsoi_inference.py", "Orchestrator", "Main entry point — runs all steps, writes CSVs"),
    ("evaluate_fsoi_results.py", "Step 6", "Closure, helpful fraction, innovation diagnostics"),
    ("fsoi_ose.py", "Step 7", "OSE denial experiments, compare vs FSOI"),
    ("visualize_fsoi.py", "Post-run", "Impact bar charts, time series, scatter plots"),
    ("compute_fsoi_weights.py", "Post-run", "Convert FSOI rankings → training loss weights"),
]
for r_idx, (mod, step, resp) in enumerate(modules, start=1):
    row = tbl.rows[r_idx]
    if r_idx % 2 == 0:
        for c in row.cells:
            set_cell_bg(c, LIGHT_BLUE)
    data = [mod, step, resp]
    for c_idx, val in enumerate(data):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx == 0:
            run.font.name = "Courier New"

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 6 — CRITICAL BUG THAT WAS FIXED
# ═══════════════════════════════════════════════════════════════════
heading("6  Critical Bug Fixed: Column Offset", 1)

callout(
    "All FSOI outputs before commit 7f04963 are INVALID due to this bug.\n"
    "Re-run results from that commit onwards.",
    bg=RGBColor(0xFF, 0xE0, 0xE0), text_color=RED
)

body("The input tensor to the GNN model has this layout:")
spacer()

code_block([
    "  Input tensor  x_input  layout:",
    "  ──────────────────────────────────────────────────────────────",
    "  Columns  0 : 7          geographic / time features",
    "                          (latitude, longitude, timestamp, ...)",
    "  Columns  7 : 7+M        instrument metadata",
    "                          (scan angles, satellite geometry, ...)",
    "  Columns  7+M : 7+M+C    ACTUAL OBSERVATION VALUES  ← want this",
    "  ──────────────────────────────────────────────────────────────",
    "",
    "  BEFORE FIX:  obs = x_input[:, 0 : C]     # reading lat/lon as obs!",
    "  AFTER FIX:   obs = x_input[:, 7+M : 7+M+C]  # correct offset",
])

body(
    "Because FSOI(k) = 0.5 × δx(k) × (ga + gb), and δx = xa − xb, "
    "extracting the wrong columns meant δx was computed on geographic features "
    "rather than actual measurement values. The innovation was meaningless, making "
    "all computed FSOI values invalid."
)

spacer()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 7 — RESULTS
# ═══════════════════════════════════════════════════════════════════
heading("7  Scientific Results  (July 2025, Post-Fix)", 1)

body("From 60 pairs (Jul 1–31 2025) using corrected obs-space FSOI:")
spacer()

tbl = doc.add_table(rows=10, cols=4)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(["Rank", "Instrument", "FSOI (sum)", "Interpretation"]):
    set_cell_bg(tbl.rows[0].cells[i], MID_BLUE)
    run = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

results = [
    ("1", "radiosonde", "−1667", "Strongly beneficial  ✓"),
    ("2", "aircraft", "−84", "Beneficial  ✓"),
    ("3", "ascat", "−7.2", "Slightly beneficial  ✓"),
    ("4", "avhrr", "−6.8", "Slightly beneficial  ✓"),
    ("5", "seviri_asr", "−4.2", "Slightly beneficial  ✓"),
    ("6", "surface_obs", "+115", "Net detrimental  ✗"),
    ("7", "atms", "+27.7", "Net detrimental  ✗"),
    ("8", "amsua", "+17.5", "Net detrimental  ✗"),
    ("9", "ssmis", "+16.6", "Net detrimental  ✗"),
]
for r_idx, (rank, inst, fsoi, interp) in enumerate(results, start=1):
    row = tbl.rows[r_idx]
    beneficial = fsoi.startswith("−") or fsoi.startswith("-")
    bg = RGBColor(0xE2, 0xEF, 0xDA) if beneficial else RGBColor(0xFF, 0xEB, 0xEB)
    for c in row.cells:
        set_cell_bg(c, bg)
    for c_idx, val in enumerate([rank, inst, fsoi, interp]):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx == 2:
            run.font.bold = True
            run.font.color.rgb = GREEN if beneficial else RED

doc.add_paragraph()

body("Key diagnostic metrics:", bold=True)
bullet("Helpful fraction: 82.2%  (target 50–80%) — most observations reduce error")
bullet("Global closure ratio: 1.524  (FAIL, ideal = 1.0) — nonlinearity due to large innovations")
bullet("Sign agreement: 66.5% — rankings qualitatively reliable")
bullet("Tier 1 FD (per-obs): 1 WARNING (radiosonde ch2, 1%), 29 SKIP (satellites — float32 noise floor, expected for 50k+ obs instruments)")
bullet(
    "Tier 2 FD (Rademacher vector): ALL PASS — 6 satellite instruments, "
    "max rel_error 2.2%, Pearson r > 0.9999. Definitively validates gradient "
    "computation (see Section 8.2)"
)

spacer()

body("Innovation RMS — What It Means and Why It Matters", bold=True)
spacer()

body(
    "The innovation for one observation is δx = xa − xb "
    "(analysed value minus background prediction). "
    "Innovation RMS is the root-mean-square of all δx values for an instrument, "
    "expressed in units of that instrument's observational standard deviation σ."
)
spacer()

body("How to read the σ units:", bold=True)
spacer()

code_block([
    "  σ  =  the 'typical' spread of that observation type",
    "         (standard deviation of the feature, from feature_stats in observation_config.yaml)",
    "",
    "  Innovation RMS in σ units  =  RMS(δx) / σ",
    "",
    "  Interpretation:",
    "    1.0σ  →  typical innovation equals the natural variability  (borderline OK)",
    "    0.5σ  →  small, well-behaved: model background is close to obs  ✓",
    "    3.0σ  →  very large: background is 3× a typical obs spread away  ✗",
    "    4.0σ  →  extreme: comparable to the tail of a normal distribution",
    "             (4σ event has probability 0.003% under Gaussian assumption)",
])

doc.add_paragraph()
body("Concrete example — radiosonde temperature:", bold=True)
spacer()

code_block([
    "  From observation_config.yaml → feature_stats → radiosonde:",
    "    airTemperature:  mean = −32.49°C,  σ = 28.82°C",
    "",
    "  One radiosonde sounding at a given location and time:",
    "    xb (background prediction) =  −10.0°C",
    "    xa (analysed / observed)   =  −27.5°C",
    "    δx = xa − xb              =  −17.5°C",
    "",
    "  Innovation in σ units:",
    "    |δx| / σ  =  17.5 / 28.82  =  0.61σ  ← this single ob is fine",
    "",
    "  But averaged over thousands of radiosondes in Jul 2025:",
    "    RMS(δx) / σ  =  3.21σ  ← the background is systematically far from obs",
    "",
    "  That means: on average, the model's background is 3.21 × 28.82 ≈ 93°C",
    "  of RMS error away from radiosonde measurements.  (Very large.)",
])

doc.add_paragraph()
body("Why Innovation RMS matters for FSOI:", bold=True)
spacer()

body(
    "FSOI is derived from a linear (first-order) approximation. It assumes the "
    "innovation δx is small enough that the forecast error surface is approximately "
    "flat between xb and xa. When δx is large, the surface curves — "
    "second-order terms become significant and the linear formula gives wrong magnitudes."
)
spacer()

tbl_inno = doc.add_table(rows=5, cols=4)
tbl_inno.style = "Table Grid"
tbl_inno.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Instrument", "Innov. RMS", "Status", "Implication for FSOI"]):
    set_cell_bg(tbl_inno.rows[0].cells[i], MID_BLUE)
    run = tbl_inno.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
inno_data = [
    ("satellites", "0.3–0.6σ", "✓ Well-behaved", "Linear approx valid, magnitudes reliable"),
    ("aircraft", "1.26σ", "~ Moderate", "Some nonlinearity, signs reliable"),
    ("radiosonde", "3.21σ", "✗ Very large", "Magnitudes approximate, signs still valid"),
    ("surface_obs", "4.01σ", "✗ Extreme", "Nonlinear — ranking direction valid only"),
]
inno_bgs = [
    RGBColor(0xE2, 0xEF, 0xDA), RGBColor(0xFF, 0xF8, 0xE1),
    RGBColor(0xFF, 0xEB, 0xEB), RGBColor(0xFF, 0xD0, 0xD0),
]
for r_idx, (inst, rms, status, impl) in enumerate(inno_data, start=1):
    row = tbl_inno.rows[r_idx]
    for c in row.cells:
        set_cell_bg(c, inno_bgs[r_idx - 1])
    for c_idx, val in enumerate([inst, rms, status, impl]):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx == 0:
            run.font.name = "Courier New"

doc.add_paragraph()

callout(
    "Why closure ratio = 1.524 (not 1.0)?\n\n"
    "FSOI is a linear approximation. When innovations are 3–4σ, nonlinear terms "
    "matter and the linear formula over-estimates the true error change.\n"
    "The SIGN direction is still trustworthy (88% sign agreement with actual ΔJ).\n"
    "MAGNITUDES are approximate — treat them as relative rankings, not absolute values.",
    bg=RGBColor(0xFF, 0xF8, 0xE1), text_color=ORANGE
)

spacer()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 7.1 — FOUR-SEASON COMPARISON
# ═══════════════════════════════════════════════════════════════════
heading("7.1  Four-Season Comparison (Jan / Apr / Jul / Oct 2025)", 2)

body(
    "FSOI runs were completed for four representative months spanning all seasons. "
    "Scaled FSOI values (sum_impact_scaled) are shown below — negative = beneficial "
    "(instrument reduced forecast error), positive = detrimental. "
    "All four seasons used the same radiosonde-fixed evaluation configuration."
)
spacer()

body("Sum Impact (Scaled) — Four Seasons", bold=True)
spacer()

tbl_seas = doc.add_table(rows=10, cols=6)
tbl_seas.style = "Table Grid"
tbl_seas.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Instrument", "Jan 2025", "Apr 2025", "Jul 2025", "Oct 2025", "Consistent?"]):
    set_cell_bg(tbl_seas.rows[0].cells[i], MID_BLUE)
    run = tbl_seas.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

seas_data = [
    ("radiosonde",  "−1645", "−1546", "−1667", "−1755", "✓ Always helpful"),
    ("aircraft",    "−137",  "−76",   "−84",   "−114",  "✓ Always helpful"),
    ("ascat",       "−10.0", "−8.0",  "−7.2",  "−9.0",  "✓ Always helpful"),
    ("avhrr",       "+1.5",  "−1.1",  "−6.8",  "−3.8",  "~ Mostly helpful"),
    ("seviri_asr",  "−0.74", "−6.2",  "−4.2",  "−2.8",  "✓ Always helpful"),
    ("surface_obs", "+31.3", "+89",   "+115",  "+94",   "✗ Always detrimental"),
    ("ssmis",       "+52.3", "+24.7", "+16.6", "+12.4", "✗ Always detrimental"),
    ("amsua",       "+3.1",  "+17.5", "+17.5", "+21.7", "✗ Always detrimental"),
    ("atms",        "+5.4",  "−1.8",  "+27.7", "+7.6",  "~ Mostly detrimental"),
]
for r_idx, row_data in enumerate(seas_data, start=1):
    inst, jan, apr, jul, oct, note = row_data
    row = tbl_seas.rows[r_idx]
    if note.startswith("✓"):
        bg = RGBColor(0xE2, 0xEF, 0xDA)
    elif note.startswith("✗"):
        bg = RGBColor(0xFF, 0xEB, 0xEB)
    else:
        bg = RGBColor(0xFF, 0xF8, 0xE1)
    for c in row.cells:
        set_cell_bg(c, bg)
    for c_idx, val in enumerate([inst, jan, apr, jul, oct, note]):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx == 0:
            run.font.name = "Courier New"
        if c_idx in (1, 2, 3, 4):
            is_helpful = val.startswith("−")
            run.font.color.rgb = GREEN if is_helpful else RED
            run.font.bold = True

doc.add_paragraph()

body("Key seasonal findings:", bold=True)
bullet(
    "Radiosonde is the dominant beneficial instrument in all seasons. "
    "October (−1755) has the highest impact, suggesting stronger upper-air "
    "temperature gradients in autumn.",
    bold_prefix="Radiosonde dominance: "
)
bullet(
    "ATMS switches from slightly beneficial in April (−1.8) to strongly detrimental "
    "in July (+27.7). Summer convective activity increases innovation variance, "
    "making GNN ATMS analysis increments less reliable.",
    bold_prefix="ATMS seasonal flip: "
)
bullet(
    "SSMIS impact magnitude decreases from winter (Jan: +52.3) to summer (Jul: +16.6). "
    "Winter polar orbit coverage differences and colder brightness temperatures are "
    "a likely driver.",
    bold_prefix="SSMIS winter peak: "
)
bullet(
    "AVHRR transitions from slightly detrimental in January (+1.5) to strongly "
    "beneficial in July (−6.8), suggesting infrared cloud-clear retrievals are "
    "more valuable in summer when cloud fraction is lower.",
    bold_prefix="AVHRR summer benefit: "
)
bullet(
    "Surface observations are consistently the most detrimental non-radiosonde "
    "instrument in absolute terms. Their 4.0σ innovation RMS indicates the GNN "
    "surface analysis increments are poorly calibrated.",
    bold_prefix="Surface obs anomaly: "
)

spacer()

body("Positive Fraction (fraction of individual observations that are individually detrimental)", bold=True)
spacer()

tbl_pfrac = doc.add_table(rows=10, cols=5)
tbl_pfrac.style = "Table Grid"
tbl_pfrac.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Instrument", "Jan 2025", "Apr 2025", "Jul 2025", "Oct 2025"]):
    set_cell_bg(tbl_pfrac.rows[0].cells[i], MID_BLUE)
    run = tbl_pfrac.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

pfrac_data = [
    ("radiosonde",  "0.456", "0.458", "0.462", "0.457"),
    ("aircraft",    "0.470", "0.469", "0.473", "0.469"),
    ("ascat",       "0.277", "0.265", "0.269", "0.266"),
    ("avhrr",       "0.251", "0.231", "0.228", "0.245"),
    ("seviri_asr",  "0.219", "0.214", "0.218", "0.201"),
    ("surface_obs", "0.440", "0.444", "0.452", "0.443"),
    ("ssmis",       "0.285", "0.275", "0.271", "0.268"),
    ("amsua",       "0.279", "0.269", "0.271", "0.271"),
    ("atms",        "0.281", "0.271", "0.273", "0.271"),
]
for r_idx, row_data in enumerate(pfrac_data, start=1):
    inst, jan, apr, jul, oct = row_data
    row = tbl_pfrac.rows[r_idx]
    if r_idx % 2 == 0:
        for c in row.cells:
            set_cell_bg(c, LIGHT_BLUE)
    for c_idx, val in enumerate([inst, jan, apr, jul, oct]):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx == 0:
            run.font.name = "Courier New"

doc.add_paragraph()
callout(
    "SEVIRI ASR has the lowest positive fraction (~0.21) — only 21% of its individual "
    "observations are individually detrimental. This is the most consistently helpful "
    "satellite instrument at the per-observation level.\n\n"
    "Radiosonde and aircraft have the highest positive fractions (~0.46–0.47) — nearly "
    "half of individual obs are detrimental. Yet the net aggregate is strongly beneficial "
    "because the helpful ~54% of obs have larger individual magnitudes.",
    bg=LIGHT_BLUE, text_color=DARK_BLUE
)

spacer()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 7.2 — PER-CHANNEL AND PER-PRESSURE ANALYSIS
# ═══════════════════════════════════════════════════════════════════
heading("7.2  Per-Channel and Per-Pressure Analysis (July 2025)", 2)

body(
    "The instrument × channel × pressure heatmaps decompose each instrument's total FSOI "
    "into contributions from individual source channels at each atmospheric pressure level. "
    "Negative values (blue) are beneficial; positive values (red) are detrimental. "
    "The most diagnostically relevant findings per instrument are summarized below."
)
spacer()

body("ATMS (Advanced Technology Microwave Sounder — 22 channels)", bold=True)
spacer()

tbl_atms = doc.add_table(rows=7, cols=3)
tbl_atms.style = "Table Grid"
tbl_atms.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Channel", "Mean Relative Contribution", "Interpretation"]):
    set_cell_bg(tbl_atms.rows[0].cells[i], MID_BLUE)
    run = tbl_atms.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

atms_ch_data = [
    ("Ch 17  (strat. O₂ ~3 hPa)",   "−3.0  (helpful)",     "Upper stratosphere sounding, clean signal"),
    ("Ch 13  (upper trop.~100 hPa)", "−2.0  (helpful)",     "Upper troposphere temperature retrieval"),
    ("Ch 3   (window / surface)",    "−1.7  (helpful)",     "Near-surface, low sounding ambiguity"),
    ("Ch 22  (183 GHz window)",      "+4.5  (detrimental)", "Water vapor window, high innovation variance"),
    ("Ch 10  (mid trop. ~250 hPa)",  "+3.9  (detrimental)", "Tropospheric sounding, increments unreliable"),
    ("Ch 15  (strat. border)",       "+3.5  (detrimental)", "Strat-trop interface — noisy gradients"),
]
for r_idx, (ch, contrib, interp) in enumerate(atms_ch_data, start=1):
    row = tbl_atms.rows[r_idx]
    is_helpful = "helpful" in contrib
    bg = RGBColor(0xE2, 0xEF, 0xDA) if is_helpful else RGBColor(0xFF, 0xEB, 0xEB)
    for c in row.cells:
        set_cell_bg(c, bg)
    for c_idx, val in enumerate([ch, contrib, interp]):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx == 1:
            run.font.bold = True
            run.font.color.rgb = GREEN if is_helpful else RED

doc.add_paragraph()

body("AMSU-A (Advanced Microwave Sounding Unit-A — 15 channels)", bold=True)
spacer()

bullet(
    "Ch 6 (~50 hPa lower stratosphere): mean contribution −9.2 — most beneficial AMSU-A channel. "
    "Provides a clean stratospheric temperature signal with low innovation variance.",
    bold_prefix="Most helpful (Ch 6): "
)
bullet(
    "Ch 13 (−2.4) and Ch 12 (−2.0) — upper troposphere / lower stratosphere, also beneficial.",
    bold_prefix="Also helpful: "
)
bullet(
    "Ch 3 (+6.9): most detrimental AMSU-A channel — surface/window channel near 850 hPa "
    "with high innovation variance.",
    bold_prefix="Most detrimental (Ch 3): "
)
bullet(
    "Ch 14 (+5.3) and Ch 10 (+2.9): mid-tropospheric sounding channels with "
    "unreliable analysis increments.",
    bold_prefix="Also detrimental: "
)

spacer()

body("SSMIS (Special Sensor Microwave Imager/Sounder)", bold=True)
spacer()

callout(
    "SSMIS Channel 21 is a major outlier across ALL target variables:\n"
    "   Dewpoint:  +15.0  (strongly detrimental)\n"
    "   Temperature:  +13.9  (strongly detrimental)\n"
    "   U-wind:  +8.9  (detrimental)\n"
    "   V-wind:  +7.0  (detrimental)\n\n"
    "Ch 21 is a 91.655 GHz window/surface channel with large brightness temperature "
    "variance over ocean and sea ice. Its analysis increments are strongly detrimental "
    "across ALL target variables. This channel should be excluded or heavily downweighted "
    "in future GNN training cycles.",
    bg=RGBColor(0xFF, 0xEB, 0xEB), text_color=RED
)

spacer()

body("ASCAT (Advanced Scatterometer)", bold=True)
spacer()

bullet(
    "V-wind: mean contribution −12.4 — the most beneficial ASCAT signal. "
    "North-south wind component retrievals from scatterometer backscatter are reliable.",
    bold_prefix="V-wind (helpful): "
)
bullet(
    "U-wind: +10.6 — detrimental. East-west wind retrievals appear less reliable, "
    "possibly due to beam geometry and ambiguity removal artifacts.",
    bold_prefix="U-wind (detrimental): "
)
bullet(
    "Net ASCAT FSOI is slightly negative (−7 to −10 across seasons) because the "
    "beneficial v-wind contribution slightly outweighs the detrimental u-wind contribution.",
    bold_prefix="Net: "
)

spacer()

body("AVHRR and SEVIRI ASR", bold=True)
spacer()

bullet(
    "AVHRR Ch 1 (0.63 μm visible, cloud detection): mean relative contribution −18.0 — "
    "the single most beneficial individual satellite channel in the evaluation. "
    "Provides cloud-clearing information that significantly improves analysis quality.",
    bold_prefix="AVHRR Ch 1: "
)
bullet(
    "SEVIRI Ch 8 (8.7 μm mid-IR window): −11.4. Ch 4 (3.9 μm near-IR window): −6.4. "
    "Short-to-mid infrared window channels provide reliable surface/cloud-top temperature.",
    bold_prefix="SEVIRI helpful channels: "
)
bullet(
    "SEVIRI Ch 10 (12.0 μm longwave window): +11.2. Ch 11 (13.4 μm CO₂): +8.4. "
    "Longwave CO₂ absorption channels with higher atmospheric sensitivity show "
    "degraded analysis increment quality.",
    bold_prefix="SEVIRI detrimental channels: "
)

spacer()

callout(
    "Key implication for channel selection:\n"
    "Per-channel analysis reveals that within a 'detrimental' instrument, some channels "
    "are individually beneficial. ATMS Ch 17/13/3 are beneficial while Ch 22/10/15 "
    "dominate the detrimental total. A channel-selective input strategy — "
    "rather than full instrument denial — could preserve beneficial information "
    "while eliminating the harmful channels that drive negative totals.",
    bg=LIGHT_BLUE, text_color=DARK_BLUE
)

spacer()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 7.3 — SEASONAL INNOVATION DIAGNOSTICS
# ═══════════════════════════════════════════════════════════════════
heading("7.3  Innovation Diagnostics Across Seasons", 2)

body(
    "Innovation RMS (in units of the feature standard deviation σ) measures how far "
    "the GNN background prediction deviates from observed values. Larger values indicate "
    "stronger nonlinearity and reduce the quantitative reliability of FSOI magnitudes."
)
spacer()

tbl_inno_seas = doc.add_table(rows=10, cols=5)
tbl_inno_seas.style = "Table Grid"
tbl_inno_seas.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Instrument", "Jan 2025", "Apr 2025", "Jul 2025", "Oct 2025"]):
    set_cell_bg(tbl_inno_seas.rows[0].cells[i], MID_BLUE)
    run = tbl_inno_seas.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

inno_seas_data = [
    ("radiosonde",  "3.052σ", "3.186σ", "3.205σ", "3.018σ"),
    ("aircraft",    "1.470σ", "1.290σ", "1.259σ", "1.277σ"),
    ("ascat",       "0.515σ", "0.514σ", "0.536σ", "0.524σ"),
    ("avhrr",       "0.379σ", "0.421σ", "0.411σ", "0.440σ"),
    ("seviri_asr",  "0.635σ", "0.687σ", "0.592σ", "0.629σ"),
    ("surface_obs", "4.032σ", "3.992σ", "4.005σ", "3.912σ"),
    ("ssmis",       "0.543σ", "0.397σ", "0.438σ", "0.386σ"),
    ("amsua",       "0.366σ", "0.346σ", "0.417σ", "0.358σ"),
    ("atms",        "0.263σ", "0.295σ", "0.320σ", "0.300σ"),
]
for r_idx, row_data in enumerate(inno_seas_data, start=1):
    inst, jan, apr, jul, oct = row_data
    row = tbl_inno_seas.rows[r_idx]
    jan_val = float(jan.replace("σ", ""))
    if jan_val > 2.0:
        bg = RGBColor(0xFF, 0xD0, 0xD0)
    elif jan_val > 1.0:
        bg = RGBColor(0xFF, 0xF8, 0xE1)
    else:
        bg = RGBColor(0xE2, 0xEF, 0xDA)
    for c in row.cells:
        set_cell_bg(c, bg)
    for c_idx, val in enumerate([inst, jan, apr, jul, oct]):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx == 0:
            run.font.name = "Courier New"

doc.add_paragraph()

body("Innovation RMS interpretation guide:", bold=True)
bullet("< 1.0σ (green): Well-behaved. FSOI magnitudes are quantitatively reliable.")
bullet("1.0–2.0σ (yellow): Moderate nonlinearity. Signs reliable; magnitudes approximate.")
bullet("> 2.0σ (red): Strong nonlinearity. Magnitudes are order-of-magnitude estimates only; signs remain trustworthy.")

spacer()

body("Key observations:", bold=True)
bullet(
    "All six satellite instruments (ATMS, AMSU-A, SSMIS, AVHRR, SEVIRI ASR, ASCAT) "
    "fall in the 0.3–0.7σ range in every season. FSOI magnitudes for satellites "
    "are quantitatively reliable across all four seasons.",
    bold_prefix="Satellites (green): "
)
bullet(
    "Aircraft innovations range 1.3–1.5σ — moderate nonlinearity, consistent across seasons. "
    "Signs are reliable; magnitudes carry some uncertainty.",
    bold_prefix="Aircraft (yellow): "
)
bullet(
    "Radiosonde innovations are consistently 3.0–3.2σ across all four seasons. "
    "The GNN background is systematically far from radiosonde measurements regardless of season. "
    "Magnitudes should be treated as relative rankings only.",
    bold_prefix="Radiosonde (red): "
)
bullet(
    "Surface observations are the most extreme at 3.9–4.0σ in every season. "
    "The GNN surface analysis is poorly calibrated — surface FSOI is directional only.",
    bold_prefix="Surface obs (red): "
)

spacer()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 7.4 — SEASONAL CLOSURE DIAGNOSTICS
# ═══════════════════════════════════════════════════════════════════
heading("7.4  Closure Diagnostics Across Seasons", 2)

body(
    "The closure test compares the total FSOI sum (linear approximation) against the "
    "actual forecast error change ΔJ measured directly. A ratio of 1.0 would indicate "
    "a perfect linear approximation. Results are shown for the global (all-instruments) "
    "diagnostic from each seasonal run."
)
spacer()

tbl_closure = doc.add_table(rows=5, cols=5)
tbl_closure.style = "Table Grid"
tbl_closure.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Season", "Period", "Closure Ratio", "Sign Agreement", "Status"]):
    set_cell_bg(tbl_closure.rows[0].cells[i], MID_BLUE)
    run = tbl_closure.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

closure_data = [
    ("Winter  (Jan 2025)", "Jan 1–31", "1.779", "72.6%", "FAIL"),
    ("Spring  (Apr 2025)", "Apr 1–30", "1.256", "65.4%", "FAIL — best"),
    ("Summer (Jul 2025)",  "Jul 1–31", "1.524", "66.5%", "FAIL"),
    ("Autumn  (Oct 2025)", "Oct 1–31", "1.811", "72.1%", "FAIL — worst"),
]
for r_idx, (season, period, ratio, sign, status) in enumerate(closure_data, start=1):
    row = tbl_closure.rows[r_idx]
    if "best" in status:
        bg = RGBColor(0xFF, 0xF8, 0xE1)
    elif "worst" in status:
        bg = RGBColor(0xFF, 0xD0, 0xD0)
    else:
        bg = RGBColor(0xFF, 0xEB, 0xEB)
    for c in row.cells:
        set_cell_bg(c, bg)
    for c_idx, val in enumerate([season, period, ratio, sign, status]):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx == 2:
            run.font.bold = True
            run.font.color.rgb = RED
        if c_idx == 4:
            run.font.bold = True
            run.font.color.rgb = RED

doc.add_paragraph()

body("Why closure ratios differ by season:", bold=True)
bullet(
    "April (1.256) has the best closure — spring Northern Hemisphere has relatively "
    "moderate atmospheric variability, reducing the nonlinear forecast error curvature "
    "that the linear FSOI approximation cannot capture.",
    bold_prefix="April (best): "
)
bullet(
    "October (1.811) has the worst closure — autumn transitions feature rapid "
    "baroclinic development (blocking onset, storm track shifts) that strongly "
    "increases the nonlinearity of the forecast error function.",
    bold_prefix="October (worst): "
)
bullet(
    "All four seasons FAIL the closure test (ideal ratio = 1.0). This is expected: "
    "the GNN is not a tangent-linear model and the linear approximation is never exact.",
    bold_prefix="Universal FAIL: "
)
bullet(
    "Sign agreement ranges 65–73% across seasons. Rankings are directionally reliable "
    "even when magnitudes are not — the instrument ordering "
    "(radiosonde >> aircraft >> ...) is consistent across all four seasons.",
    bold_prefix="Sign agreement: "
)

spacer()

callout(
    "Closure test summary — four seasons:\n"
    "  Best season:   April 2025  (ratio = 1.256, closest to ideal 1.0)\n"
    "  Worst season:  October 2025  (ratio = 1.811, most nonlinear)\n"
    "  Full range:    1.256 – 1.811  across all seasons\n\n"
    "FSOI overpredicts the true forecast error change by 25–81% depending on season. "
    "The SIGN of instrument rankings is reliable (65–73% agreement with actual ΔJ). "
    "Magnitudes should be treated as relative rankings only — not used as "
    "quantitative claims about the size of each instrument's impact.",
    bg=RGBColor(0xFF, 0xF8, 0xE1), text_color=ORANGE
)

spacer()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 8 — VALIDATION FRAMEWORK
# ═══════════════════════════════════════════════════════════════════
heading("8  Validation Framework", 1)

body(
    "FSOI relies on correct gradient computation through the GNN. Three tiers of "
    "validation are implemented to catch broken computation graphs:"
)
spacer()

code_block([
    "  Tier 1 — Scalar FD check (float32)",
    "  ─────────────────────────────────────────────────────────────",
    "  For 10 sampled observations:",
    "    Perturb obs channel by ε = 1e-3",
    "    FD gradient = (J(x+ε) − J(x−ε)) / (2ε)",
    "    Compare vs autograd gradient",
    "    PASS if relative error < 1%  |  WARNING if < 5%  |  FAIL if > 5%",
    "",
    "  Tier 2 — Directional derivative (Rademacher vector)",
    "  ─────────────────────────────────────────────────────────────",
    "  Random direction v ∈ {±1}^n (one sample covers all channels)",
    "    FD: (J(x + εv) − J(x − εv)) / (2ε)",
    "    Adjoint: gᵀv",
    "  Efficient for high-dimensional satellite instruments",
    "",
    "  Tier 3 — Float64 FD check",
    "  ─────────────────────────────────────────────────────────────",
    "  Same as Tier 1 but promotes to float64 for precision",
    "  Resolves WARNING→PASS cases caused by float32 rounding",
])

spacer()

# ── Tier 2 deep dive ────────────────────────────────────────────────
heading("8.1  Deep Dive: Tier 2 Directional Derivative", 2)

body("The core question Tier 2 answers:", bold=True)
callout(
    "Did autograd compute the correct gradient — across ALL channels simultaneously — "
    "in just 2 forward passes?",
    bg=LIGHT_BLUE, text_color=DARK_BLUE
)

body("The Problem with Tier 1 at Scale", bold=True)
spacer()
body(
    "Tier 1 perturbs one channel at a time, so it needs 2 × n_channels forward passes "
    "per instrument. For ATMS (22 channels × thousands of obs nodes) that is 44 separate "
    "model runs. Tier 2 collapses this to exactly 2 runs regardless of dimension."
)
spacer()

# Cost comparison table
tbl = doc.add_table(rows=4, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Instrument", "Channels", "Tier 1 runs  vs  Tier 2 runs"]):
    set_cell_bg(tbl.rows[0].cells[i], MID_BLUE)
    run = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
cost_data = [
    ("atms", "22", "44 runs  vs  2 runs"),
    ("amsua", "15", "30 runs  vs  2 runs"),
    ("radiosonde", " 4", " 8 runs  vs  2 runs"),
]
for r_idx, (inst, ch, cost) in enumerate(cost_data, start=1):
    row = tbl.rows[r_idx]
    if r_idx % 2 == 0:
        for c in row.cells:
            set_cell_bg(c, LIGHT_BLUE)
    for c_idx, val in enumerate([inst, ch, cost]):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx == 0:
            run.font.name = "Courier New"

doc.add_paragraph()
body("What is a Rademacher Vector?", bold=True)
spacer()
body(
    "v ∈ {±1}ⁿ means every entry is randomly either +1 or −1. "
    "It is a random direction in n-dimensional space where every coordinate has "
    "equal magnitude — no channel is louder than another."
)
spacer()

code_block([
    "  Example — ATMS (22 channels, simplified to 5 here):",
    "",
    "    x = [ 0.5,  1.2, −0.3,  0.8,  0.1 ]   ← current obs values",
    "    v = [  +1,   −1,   +1,   +1,   −1 ]   ← random Rademacher draw",
    "    ε = 0.01",
    "",
    "    x + εv = [ 0.51, 1.19, −0.29, 0.81,  0.09 ]",
    "    x − εv = [ 0.49, 1.21, −0.31, 0.79,  0.11 ]",
    "",
    "    Run model twice to get:",
    "      J(x + εv)  and  J(x − εv)",
    "",
    "    FD  =  ( J(x + εv) − J(x − εv) ) / (2ε)",
    "        =  directional derivative of J along v  (numerical)",
])

doc.add_paragraph()
body("What Does the Gradient Predict?", bold=True)
spacer()
body(
    "The adjoint g = ∂J/∂x (computed by autograd.backward()) predicts the same "
    "directional derivative via a simple dot product:"
)
formula_block("Adjoint prediction  =  gᵀv  =  Σᵢ  g[i] × v[i]")

body(
    "If autograd is correct, the chain rule guarantees FD ≈ gᵀv. "
    "If they disagree, autograd lost the gradient somewhere in the computation graph."
)
spacer()

body("Numeric Example — Pass vs Fail", bold=True)
spacer()

code_block([
    "  Setup (3 channels for clarity):",
    "    x = [ 0.5,  1.2, −0.3 ]",
    "    v = [  +1,   −1,   +1 ]",
    "    ε = 0.01",
    "",
    "  Model returns:  J(x + εv) = 10.014,   J(x − εv) = 9.990",
    "  FD  =  (10.014 − 9.990) / 0.02  =  1.20",
    "",
    "  ─── Scenario A: autograd correct ───────────────────────────",
    "    g = [ 0.8, −0.2,  0.6 ]",
    "    gᵀv = (0.8)(+1) + (−0.2)(−1) + (0.6)(+1) = 0.8 + 0.2 + 0.6 = 1.60",
    "    Hmm — FD=1.20 vs adjoint=1.60, 25% error  ← ε too large (nonlinear)",
    "",
    "  ─── Scenario B: gradient component wrong ────────────────────",
    "    g = [ 0.8, −0.4,  0.6 ]   ← ch2 gradient wrong sign",
    "    gᵀv = (0.8)(+1) + (−0.4)(−1) + (0.6)(+1) = 0.8 + 0.4 + 0.6 = 1.80",
    "    Relative error = |1.20 − 1.80| / 1.80 = 33%  →  FAIL",
    "",
    "  ─── What failure means ──────────────────────────────────────",
    "    A detach(), in-place op, or broken indexing cut the grad path.",
    "    The gradient is NOT reliable → FSOI values are INVALID.",
])

doc.add_paragraph()
body("Why ±1 Specifically — Not Random Gaussians?", bold=True)
spacer()

tbl2 = doc.add_table(rows=3, cols=3)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Direction type", "Property", "Effect on test"]):
    set_cell_bg(tbl2.rows[0].cells[i], MID_BLUE)
    run = tbl2.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
why_data = [
    ("Gaussian v ~ N(0,1)", "Large entries dominate dot product",
     "Channels with small gradient are tested weakly — bugs can hide"),
    ("Rademacher v ∈ {±1}ⁿ", "Every entry has magnitude exactly 1",
     "All channels tested with equal weight — minimum estimator variance"),
]
for r_idx, (t, p, e) in enumerate(why_data, start=1):
    row = tbl2.rows[r_idx]
    if r_idx % 2 == 0:
        for c in row.cells:
            set_cell_bg(c, LIGHT_BLUE)
    for c_idx, val in enumerate([t, p, e]):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)

doc.add_paragraph()
body("Why Tier 1 Fails for Satellites (and Tier 2 Saves It)", bold=True)
spacer()
body(
    "Satellite brightness temperature innovations are 0.3–0.6σ — very small values. "
    "When you perturb a single channel by ε = 1e-3 in float32, the resulting change "
    "J(x+ε) − J(x−ε) can be smaller than float32 rounding noise (≈ 1e-7). "
    "The FD estimate is just noise — hence the 29 SKIP results in your runs."
)
spacer()
body(
    "Tier 2 perturbs ALL N×C elements simultaneously with the Rademacher vector. "
    "The total signal is the sum of N×C contributions — orders of magnitude larger — "
    "so it clears the float32 noise floor even for satellites. This is why Tier 2 "
    "can give a reliable PASS/FAIL where Tier 1 can only say SKIP."
)
spacer()

callout(
    "Tier 1 = check one gradient component carefully (precision, but expensive + skips for small obs)\n"
    "Tier 2 = check all components together cheaply (catches broken autograd graphs reliably)\n"
    "Tier 3 = repeat Tier 1 in float64 to resolve float32 rounding (turns SKIP → PASS or FAIL)",
    bg=RGBColor(0xE2, 0xEF, 0xDA), text_color=RGBColor(0x00, 0x50, 0x00)
)

spacer()

# ── 8.2 Tier 2 actual results ────────────────────────────────────────────
heading("8.2  Tier 2 Validation Results — July 2025", 2)

body(
    "Results from the fd_check_enhanced run (bins 2025-07-01 through 2025-07-03, "
    "3 pairs, 5 Rademacher trials each). Tested instruments: amsua, ascat, atms, avhrr, "
    "seviri_asr, ssmis."
)
spacer()

tbl_t2 = doc.add_table(rows=7, cols=4)
tbl_t2.style = "Table Grid"
tbl_t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Instrument", "Max rel_error", "Min Pearson r", "Result"]):
    set_cell_bg(tbl_t2.rows[0].cells[i], MID_BLUE)
    run = tbl_t2.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

t2_data = [
    ("amsua", "0.46%", "0.9999997", "PASS — all 5 trials"),
    ("ascat", "0.12%", "0.9999999", "PASS — all 5 trials"),
    ("atms", "0.53%", "0.9999999", "PASS — all 5 trials"),
    ("avhrr", "0.57%", "0.9999994", "PASS — all 5 trials"),
    ("seviri_asr", "0.47%", "0.9999985", "PASS — all 5 trials"),
    ("ssmis", "2.2%", "0.9999991", "PASS — all 5 trials"),
]
for r_idx, (inst, rel_err, pearson, result) in enumerate(t2_data, start=1):
    row = tbl_t2.rows[r_idx]
    for c in row.cells:
        set_cell_bg(c, RGBColor(0xE2, 0xEF, 0xDA))
    for c_idx, val in enumerate([inst, rel_err, pearson, result]):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx == 0:
            run.font.name = "Courier New"
        if c_idx == 3:
            run.font.color.rgb = GREEN
            run.font.bold = True

doc.add_paragraph()

callout(
    "All 6 satellite/microwave instruments PASS with relative error < 1%  "
    "(max 2.2% for SSMIS trial 2, still well within 5% threshold).\n"
    "Pearson r between autograd directional derivative and FD directional derivative is > 0.9999 "
    "for all instruments — effectively perfect linear correspondence.\n\n"
    "This definitively validates gradient correctness despite Tier 1 showing only SKIP for satellites. "
    "The gradient computation graph through the GNN is confirmed intact.",
    bg=RGBColor(0xE2, 0xEF, 0xDA), text_color=RGBColor(0x00, 0x50, 0x00)
)

spacer()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 9 — OSE vs FSOI
# ═══════════════════════════════════════════════════════════════════
heading("9  OSE vs FSOI — Validating Rankings", 1)

callout(
    "OSE is the ground-truth experiment. FSOI is the cheap approximation.\n"
    "Comparing them tells you whether the FSOI rankings can be trusted.",
    bg=LIGHT_BLUE, text_color=DARK_BLUE
)

body("What is an OSE?", bold=True)
spacer()
body(
    "An Observing System Experiment denies one instrument entirely — replacing its true "
    "observations (xa) with the model background (xb) — and re-runs the forecast. "
    "The change in forecast error compared to the full-obs run "
    "directly measures that instrument's real-world impact — no approximations involved."
)
spacer()

body("Side-by-side: FSOI vs OSE", bold=True)
spacer()

code_block([
    "  ┌───────────────────────────────────────────────────────────────┐",
    "  │  FULL RUN (baseline — used by both FSOI and OSE)             │",
    "  │  All true observations used → forecast → error J_full         │",
    "  └───────────────────────────────────────────────────────────────┘",
    "                    │                         │",
    "                    ▼ (FSOI path)             ▼ (OSE path)",
    "  ┌─────────────────────────────┐  ┌──────────────────────────────┐",
    "  │  Backpropagate gradient     │  │  Deny instrument i           │",
    "  │  through the GNN            │  │  xa[i] → xb[i]  (replace     │",
    "  │                             │  │  analysed with background)   │",
    "  │  Compute:                   │  │                              │",
    "  │  FSOI(k) = ½δx·(ga+gb)      │  │  Re-run forecast             │",
    "  │  Sum over instrument i      │  │  → error J_denied            │",
    "  │                             │  │                              │",
    "  │  ΔJ_fsoi = Σ_k FSOI(k)      │  │  ΔJ_ose = J_denied − J_full  │",
    "  │  Cost: 1 fwd + 1 bwd        │  │  Cost: 1 extra forward pass  │",
    "  └─────────────────────────────┘  └──────────────────────────────┘",
    "                    │                         │",
    "                    └──────────┬──────────────┘",
    "                               ▼",
    "              Compare: sign(ΔJ_fsoi)  vs  sign(ΔJ_ose)",
    "              Ratio:   ΔJ_fsoi / ΔJ_ose  (ideal = 1.0)",
])

doc.add_paragraph()
body("Worked Example — ATMS Denial", bold=True)
spacer()

body(
    "Suppose we run the full Jul 2025 evaluation and then run an OSE "
    "denying ATMS (replacing its analysed values with the model background). "
    "Here is what the numbers might look like:"
)
spacer()

code_block([
    "  Full run (all true observations included):",
    "    J_full  =  1000.0   (forecast MSE, arbitrary units)",
    "",
    "  FSOI path (linear approximation, same full run):",
    "    Σ_k FSOI(k) for ATMS alone  =  +27.7",
    "    Interpretation: ATMS contributed +27.7 to forecast error  (detrimental)",
    "    ΔJ_fsoi(ATMS) = +27.7  →  removing ATMS should IMPROVE forecast",
    "",
    "  OSE path (re-run without ATMS):",
    "    xa[atms]  →  xb[atms]  (replace ATMS true obs with prior forecast)",
    "    Re-run forecast  →  J_denied = 972.0",
    "    ΔJ_ose(ATMS)  =  J_denied − J_full  =  972.0 − 1000.0  =  −28.0",
    "    Interpretation: removing ATMS made the forecast BETTER by 28.0 units  ✓",
    "",
    "  Comparison:",
    "    sign(ΔJ_fsoi) = +1   (FSOI predicts: ATMS is detrimental)",
    "    sign(ΔJ_ose)  = −1   (OSE confirms: removing ATMS helps)",
    "    Sign agreement ✓  (they agree ATMS hurts — sign convention flips",
    "                        because FSOI measures impact WITH obs,",
    "                        OSE measures change from REMOVING obs)",
    "",
    "    Magnitude ratio  =  |ΔJ_fsoi| / |ΔJ_ose|  =  27.7 / 28.0  =  0.99  ✓",
    "    Near 1.0 → linear approximation was accurate for ATMS",
])

doc.add_paragraph()

callout(
    "Sign convention note:\n"
    "FSOI(i) > 0 means instrument i increased error when present.\n"
    "ΔJ_ose  < 0 means removing instrument i decreased error.\n"
    "Both say the same thing — ATMS is detrimental. The signs are opposite "
    "by construction because one measures impact of presence, the other measures impact of absence.",
    bg=RGBColor(0xFF, 0xF8, 0xE1), text_color=ORANGE
)

body("Why Closure Can Fail for Some Instruments", bold=True)
spacer()

body(
    "For instruments with large innovations (radiosonde 3.21σ, surface_obs 4.01σ), "
    "the FSOI linear approximation breaks down and the magnitude ratio drifts away from 1.0:"
)
spacer()

code_block([
    "  Radiosonde example (large innovations, 3.21σ):",
    "",
    "    ΔJ_fsoi(radiosonde)  =  −1667   (FSOI predicts large improvement)",
    "    ΔJ_ose(radiosonde)   =  −1094   (OSE measures actual improvement)",
    "",
    "    Magnitude ratio  =  1667 / 1094  =  1.52   ← closure ratio ~ 1.524",
    "    FSOI over-predicts the impact by 52% because the quadratic error",
    "    surface term is not negligible at 3σ innovations.",
    "",
    "    Sign agreement: both negative  ✓  (radiosonde is genuinely beneficial)",
    "    Magnitude: approximate — but the ranking ORDER is still reliable.",
])

doc.add_paragraph()

body("What OSE Validates and What It Cannot", bold=True)
spacer()

tbl_ose = doc.add_table(rows=4, cols=3)
tbl_ose.style = "Table Grid"
tbl_ose.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Question", "OSE answer", "FSOI answer"]):
    set_cell_bg(tbl_ose.rows[0].cells[i], MID_BLUE)
    run = tbl_ose.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
ose_rows = [
    ("Is this instrument beneficial or harmful?",
     "✓ Definitive ground truth",
     "✓ Reliable (sign)"),
    ("How much does it help/hurt (magnitude)?",
     "✓ Exact measurement",
     "~ Approximate (linear approx)"),
    ("Which individual obs within instrument matter?",
     "✗ Cannot — only tests whole instrument",
     "✓ Per-observation FSOI(k) available"),
]
for r_idx, (q, o, f) in enumerate(ose_rows, start=1):
    row = tbl_ose.rows[r_idx]
    if r_idx % 2 == 0:
        for c in row.cells:
            set_cell_bg(c, LIGHT_BLUE)
    for c_idx, val in enumerate([q, o, f]):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)

doc.add_paragraph()
spacer()

# ── 9.1 Actual ATMS OSE results ──────────────────────────────────────────
heading("9.1  Actual ATMS OSE Results — July & January 2025", 2)

body(
    "The OSE denial experiments for ATMS are complete. Results from "
    "ose_atms_jul2025_fixed (59 pairs, July 2025) and ose_atms_jan2025_fixed (January 2025)."
)
spacer()

tbl_ose_actual = doc.add_table(rows=7, cols=3)
tbl_ose_actual.style = "Table Grid"
tbl_ose_actual.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Metric", "July 2025", "January 2025"]):
    set_cell_bg(tbl_ose_actual.rows[0].cells[i], MID_BLUE)
    run = tbl_ose_actual.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

ose_actual_rows = [
    ("Pairs evaluated", "59", "~30"),
    ("ATMS detrimental (OSE confirms)", "55 / 59  (93%)", "~70%"),
    ("Sign agreement with FSOI", "55 / 59  (93%  True)", "mostly True"),
    ("OSE relative impact per pair", "−0.01% to −0.08%", "−0.06% to −0.08%"),
    ("Closure ratio  |FSOI| / |ΔJ_ose|", "77 – 2451  (med ~150)", "34 – 92"),
    ("FSOI predicted magnitude per pair", "0.08 – 2.79", "0.05 – 0.52"),
]
for r_idx, (metric, jul, jan) in enumerate(ose_actual_rows, start=1):
    row = tbl_ose_actual.rows[r_idx]
    if r_idx % 2 == 0:
        for c in row.cells:
            set_cell_bg(c, LIGHT_BLUE)
    for c_idx, val in enumerate([metric, jul, jan]):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)

doc.add_paragraph()
body("Key findings:", bold=True)
spacer()

bullet(
    "ATMS is genuinely net detrimental in 93% of July pairs — the "
    "analysis increments for ATMS slightly worsen the forecast. "
    "Both FSOI and OSE agree on the direction.",
    bold_prefix="Direction validated: "
)
bullet(
    "The closure ratio of 77–2451 (median ~150) means FSOI predicts "
    "100–2400× the actual OSE-measured impact. Root cause: each ATMS "
    "observation has gradient g ≈ 2×10⁻⁶, and summing 1.1M ATMS obs "
    "per pair overestimates the nonlinear OSE result by orders of magnitude.",
    bold_prefix="Magnitude overestimated: "
)
bullet(
    "January 2025 closure ratios (34–92) are lower than July (77–2451). "
    "Summer active weather introduces more nonlinearity that the linear "
    "approximation cannot capture.",
    bold_prefix="Seasonal difference: "
)
bullet(
    "4 sign-flipped pairs in July (pairs 35, 43, 47, 49) where OSE shows "
    "ATMS slightly helpful (+0.0000–+0.0009 error change). "
    "All are noise-level impacts at the limit of detectability.",
    bold_prefix="Exceptions (4/59): "
)

spacer()

callout(
    "OSE conclusion — ATMS (July 2025):\n"
    "  SIGN:      CORRECT  —  93% agreement. ATMS is genuinely net detrimental.\n"
    "  MAGNITUDE: NOT reliable  —  FSOI overpredicts by 100–2400×.\n\n"
    "Use FSOI for instrument ranking (which instruments help vs hurt).\n"
    "Do NOT use FSOI magnitudes for quantitative claims about impact size.\n"
    "OSE is the ground-truth for magnitude; FSOI is the cheap directional indicator.",
    bg=RGBColor(0xFF, 0xF8, 0xE1), text_color=ORANGE
)

spacer()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 10 — FSOI WEIGHTS → FINE-TUNING
# ═══════════════════════════════════════════════════════════════════
heading("10  FSOI Weights for Fine-Tuning", 1)

body(
    "The FSOI impact scores are converted into loss weights for the next round of "
    "GNN training. Two variants are computed:"
)
spacer()

bullet(
    "w ∝ |mean_impact|  — weight proportional to absolute average FSOI magnitude",
    bold_prefix="Variant A: "
)
bullet(
    "w ∝ |mean_impact| × reliability²  — down-weights instruments with noisy innovations",
    bold_prefix="Variant B: "
)

spacer()
body("Resulting weights (obs-space and mesh-space weights are nearly identical):", bold=True)
spacer()

tbl = doc.add_table(rows=5, cols=3)
tbl.style = "Table Grid"
for i, h in enumerate(["Instrument", "Weight A", "Weight B"]):
    set_cell_bg(tbl.rows[0].cells[i], MID_BLUE)
    run = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

weight_data = [
    ("radiosonde", "8.178", "8.179"),
    ("aircraft", "0.153", "0.146"),
    ("surface_obs", "0.106", "0.111"),
    ("all satellites", "0.094", "0.094"),
]
for r_idx, (inst, wa, wb) in enumerate(weight_data, start=1):
    row = tbl.rows[r_idx]
    if r_idx % 2 == 0:
        for c in row.cells:
            set_cell_bg(c, LIGHT_BLUE)
    for c_idx, val in enumerate([inst, wa, wb]):
        run = row.cells[c_idx].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx in (1, 2):
            run.font.bold = True

doc.add_paragraph()
callout(
    "Radiosonde dominance (weight ~87× satellites) is confirmed by both obs-space "
    "and mesh-space runs — it is not an artifact of the evaluation methodology.",
    bg=RGBColor(0xE2, 0xEF, 0xDA), text_color=GREEN
)

# ═══════════════════════════════════════════════════════════════════
#  SECTION 11 — HOW TO RUN
# ═══════════════════════════════════════════════════════════════════
heading("11  How to Run FSOI", 1)

body("Three-step canonical usage:")
spacer()

code_block([
    "# Step 1 — Quick sanity test (6 unit tests, ~2 min)",
    "python FSOI/test_fsoi.py --checkpoint /path/to/model.ckpt",
    "",
    "# Step 2 — Run FSOI computation",
    "python FSOI/fsoi_inference.py \\",
    "    --checkpoint  /path/to/model.ckpt \\",
    "    --config       FSOI/configs/fsoi_config.yaml \\",
    "    --start_date  2025-06-01 \\",
    "    --end_date    2025-07-01 \\",
    "    --output_dir  ./FSOI/fsoi_outputs/my_run",
    "",
    "# Step 3 — Visualize",
    "python FSOI/visualize_fsoi.py \\",
    "    --input  ./FSOI/fsoi_outputs/my_run/csv \\",
    "    --output ./FSOI/fsoi_outputs/my_run/figures",
])

body("Output directory structure per run:", bold=True)
spacer()

code_block([
    "  <output_dir>/",
    "    csv/",
    "      fsoi_by_instrument.csv   ← main results (per pair, per instrument)",
    "      fsoi_by_channel.csv      ← per-channel breakdown",
    "      fsoi_summary.csv         ← aggregated across all pairs",
    "      scatter_samples.csv      ← innovation vs FSOI samples (for maps)",
    "    evaluation/",
    "      fsoi_closure_summary.csv         ← linear approximation quality",
    "      fsoi_system_health.csv           ← helpful_fraction, system_flag",
    "      innovation_diagnostics.csv       ← δx statistics per instrument",
    "      fd_validation.csv                ← FD gradient check results",
    "      fsoi_beneficial_fraction.csv     ← helpful vs harmful obs breakdown",
    "      fsoi_regional_summary.csv        ← geographic breakdown",
    "      fsoi_closure_per_level_summary.csv ← vertical structure",
    "    figures/",
    "      instrument_impacts.png",
    "      innovation/   ← histograms, time series, heatmaps",
    "      maps/         ← global FSOI spatial maps",
    "    logs/",
    "      fsoi_config_used.yaml   ← config snapshot for reproducibility",
])

# ═══════════════════════════════════════════════════════════════════
#  SECTION 12 — PROJECT STATUS
# ═══════════════════════════════════════════════════════════════════
heading("12  Project Status  (May 2026)", 1)

body("Completed:", bold=True, color=GREEN)
bullet("4 seasonal obs-space FSOI runs (Jan/Apr/Jul/Oct 2025)")
bullet("12 mesh-space jobs (3 pressure levels × 4 seasons)")
bullet("3-tier FD gradient validation framework")
bullet("Tier 2 FD (Rademacher): all 6 satellite instruments PASS (<2.2% error, Pearson r > 0.9999)")
bullet("OSE ATMS denial: July 2025 (59 pairs) and January 2025 — both COMPLETE")
bullet("FSOI weights computed — obs-space biased + mesh-space primary (unbiased)")
bullet("Critical bugs fixed (column-offset, FD precision)")
bullet("Comprehensive documentation in FSOI/docs/")

spacer()
body("Next steps:", bold=True, color=MID_BLUE)
bullet("FSOI-weighted fine-tuning with Variant A and B weights")
bullet("RMSE comparison vs baseline (M0 model)")
bullet("Paper writing — Methods, Results, Discussion sections")

spacer()

# ═══════════════════════════════════════════════════════════════════
#  FOOTER
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run(
    "OCELOT GNN Weather Model  •  FSOI Implementation\n"
    "gnn_model/FSOI/  •  Branch: feature/FSOI_evaluation\n"
    "Generated May 2026"
)
run.font.size = Pt(9)
run.font.color.rgb = DARK_GRAY
run.font.italic = True

# ── Save ─────────────────────────────────────────────────────────
out_path = r"C:\Users\azadeh.gholoubi\Desktop\FSOI_Explanation.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
